import os
import logging
from fastapi.responses import FileResponse
from typing import List, Optional, Dict
from fastapi import APIRouter, Depends, HTTPException, UploadFile, File, Form, BackgroundTasks
from fastapi.responses import JSONResponse
from sqlalchemy import func, text
from sqlalchemy.orm import Session

from app.core.database import get_db
from app.models.database_models import Document as Document, DocumentChunk
from app.services.document_processor import DocumentProcessor
from app.services.chroma_service import ChromaService
from app.config.settings import settings

logger = logging.getLogger(__name__)
router = APIRouter()

# Initialize services
document_processor = DocumentProcessor()
chroma_service = ChromaService()


def validate_file_type(filename: str) -> bool:
    """Validate if file type is supported"""
    file_extension = filename.lower().split('.')[-1]
    return file_extension in settings.allowed_file_types


def validate_file_size(file_size: int) -> bool:
    """Validate file size"""
    return file_size <= settings.max_file_size
async def complete_document_processing(document_id: int, chunks: List[Dict], metadata: Dict):
    """Complete document processing by saving chunks and adding to vector store"""
    from app.core.database import SessionLocal

    db = None
    try:
        # Create database session
        db = SessionLocal()

        # Get document
        document = db.query(Document).filter(Document.id == document_id).first()
        if not document:
            logger.error(f"Document {document_id} not found for processing completion")
            return

        logger.info(f"Completing processing for document {document_id}: {document.filename}")

        # Save chunks to database
        for chunk_data in chunks:
            chunk = DocumentChunk(
                document_id=document_id,
                chunk_index=chunk_data["chunk_index"],
                content=chunk_data["content"],
                word_count=chunk_data["word_count"],
                char_count=chunk_data["char_count"],
                page_number=chunk_data.get("page_number")
            )
            db.add(chunk)

        # Initialize ChromaDB if not already done
        await chroma_service.initialize()

        # Add to vector store
        documents_for_vector = []
        metadatas_for_vector = []
        ids_for_vector = []

        for i, chunk_data in enumerate(chunks):
            documents_for_vector.append(chunk_data["content"])
            metadatas_for_vector.append({
                "document_id": document_id,
                "chunk_index": i,
                "file_type": document.file_type,
                "filename": document.filename,  # Use the final filename
                "original_filename": document.original_filename,
                "title": document.title or "",
                "category": document.category or "",
                "page_number": chunk_data.get("page_number", 0)
            })
            ids_for_vector.append(f"doc_{document_id}_chunk_{i}")

        # Add to ChromaDB
        success = await chroma_service.add_documents(
            documents=documents_for_vector,
            metadatas=metadatas_for_vector,
            ids=ids_for_vector
        )

        if success:
            # Update document status
            document.status = "completed"
            document.processing_status = "completed"
            document.is_processed = True
            document.is_searchable = True
            document.processed_at = func.now()
            document.error_message = None

            logger.info(f"Document {document_id} processing completed successfully")
        else:
            raise Exception("Failed to add documents to vector store")

        db.commit()

    except Exception as e:
        logger.error(f"Error completing document processing {document_id}: {e}")

        if db:
            try:
                # Update document with error status
                document = db.query(Document).filter(Document.id == document_id).first()
                if document:
                    document.status = "failed"
                    document.processing_status = "failed"
                    document.error_message = str(e)
                    db.commit()
            except Exception as commit_error:
                logger.error(f"Error updating document status: {commit_error}")

    finally:
        if db:
            db.close()


@router.get("/")
async def list_documents(
        skip: int = 0,
        limit: int = 100,
        file_type: Optional[str] = None,
        category: Optional[str] = None,
        status: Optional[str] = None,
        db: Session = Depends(get_db)
):
    """List all documents with optional filtering"""

    try:
        # Build WHERE clause
        where_conditions = []
        params = {}

        if file_type:
            where_conditions.append("file_type = :file_type")
            params["file_type"] = file_type
        if category:
            where_conditions.append("category = :category")
            params["category"] = category
        if status:
            where_conditions.append("status = :status")
            params["status"] = status

        where_clause = " WHERE " + " AND ".join(where_conditions) if where_conditions else ""

        # Count query
        count_sql = f"SELECT COUNT(*) as total FROM documents{where_clause}"
        count_result = db.execute(text(count_sql), params).fetchone()
        total = count_result.total if count_result else 0

        # Main query
        params.update({"skip": skip, "limit": limit})
        main_sql = f"""
        SELECT 
            id,
            filename,
            original_filename,
            file_type,
            file_size,
            title,
            category,
            description,
            status,
            processing_status,
            created_at,
            processed_at,
            word_count,
            total_chunks
        FROM documents
        {where_clause}
        ORDER BY created_at DESC
        LIMIT :limit OFFSET :skip
        """

        results = db.execute(text(main_sql), params).fetchall()

        # Convert to list of dictionaries
        document_list = []
        for row in results:
            # Calculate is_processed based on processing_status
            is_processed = row.processing_status == "completed" if row.processing_status else False

            # Handle datetime fields - they might already be strings
            created_at = row.created_at
            if hasattr(row.created_at, 'isoformat'):
                created_at = row.created_at.isoformat()

            processed_at = row.processed_at
            if hasattr(row.processed_at, 'isoformat'):
                processed_at = row.processed_at.isoformat()

            doc_dict = {
                "id": row.id,
                "filename": row.original_filename or row.filename or "Unknown",
                "file_type": row.file_type,
                "file_size": row.file_size,
                "title": row.title,
                "category": row.category,
                "description": row.description,
                "status": row.status,
                "processing_status": row.processing_status,
                "is_processed": is_processed,
                "created_at": created_at,
                "processed_at": processed_at,
                "word_count": row.word_count,
                "total_chunks": row.total_chunks,
                "metadata": {}
            }
            document_list.append(doc_dict)

        return {
            "documents": document_list,
            "total": total,
            "skip": skip,
            "limit": limit,
            "filters": {
                "file_type": file_type,
                "category": category,
                "status": status
            }
        }

    except Exception as e:
        logger.error(f"Error listing documents: {e}")
        import traceback
        logger.error(f"Full traceback: {traceback.format_exc()}")
        raise HTTPException(status_code=500, detail=f"Failed to list documents: {str(e)}")


@router.get("/{document_id}")
async def get_document(document_id: int, db: Session = Depends(get_db)):
    """Get document details by ID"""

    try:
        document = db.query(Document).filter(Document.id == document_id).first()

        if not document:
            raise HTTPException(status_code=404, detail="Document not found")

        # Get chunks count
        chunks_count = db.query(DocumentChunk).filter(DocumentChunk.document_id == document_id).count()

        return {
            "id": document.id,
            "filename": document.original_filename,
            "file_type": document.file_type,
            "file_size": document.file_size,
            "file_hash": document.file_hash,
            "title": document.title,
            "author": document.author,
            "subject": document.subject,
            "creator": document.creator,
            "category": document.category,
            "description": document.description,
            "keywords": document.keywords,
            "total_pages": document.total_pages,
            "total_chunks": document.total_chunks,
            "word_count": document.word_count,
            "char_count": document.char_count,
            "chunks_in_db": chunks_count,
            "status": document.status,
            "processing_status": document.processing_status,
            "error_message": document.error_message,
            "is_processed": document.is_processed,
            "is_searchable": document.is_searchable,
            "created_at": document.created_at.isoformat() if document.created_at else None,
            "updated_at": document.updated_at.isoformat() if document.updated_at else None,
            "processed_at": document.processed_at.isoformat() if document.processed_at else None,
            "metadata": document.metadata
        }

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error getting document {document_id}: {e}")
        raise HTTPException(status_code=500, detail=f"Failed to get document: {str(e)}")


@router.post("/{document_id}/process")
async def process_document(
        document_id: int,
        background_tasks: BackgroundTasks,
        db: Session = Depends(get_db)
):
    """Manually trigger document processing"""

    try:
        document = db.query(Document).filter(Document.id == document_id).first()

        if not document:
            raise HTTPException(status_code=404, detail="Document not found")

        if document.processing_status == "processing":
            raise HTTPException(status_code=400, detail="Document is already being processed")

        # Update status
        document.processing_status = "processing"
        db.commit()

        # Start background processing
        background_tasks.add_task(process_document_background, document_id)

        return {
            "message": "Document processing started",
            "document_id": document_id,
            "status": "processing"
        }

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error starting document processing: {e}")
        raise HTTPException(status_code=500, detail=f"Failed to start processing: {str(e)}")


@router.delete("/{document_id}")
async def delete_document(document_id: int, db: Session = Depends(get_db)):
    """Delete a document and its associated data"""

    try:
        document = db.query(Document).filter(Document.id == document_id).first()

        if not document:
            raise HTTPException(status_code=404, detail="Document not found")

        # Delete from vector store
        try:
            await chroma_service.delete_documents(document_id)
        except Exception as e:
            logger.warning(f"Error deleting from vector store: {e}")

        # Delete chunks from database
        db.query(DocumentChunk).filter(DocumentChunk.document_id == document_id).delete()

        # Delete file from filesystem
        if os.path.exists(document.file_path):
            try:
                os.remove(document.file_path)
                logger.info(f"Deleted file: {document.file_path}")
            except Exception as e:
                logger.warning(f"Could not delete file {document.file_path}: {e}")

        # Delete document record
        db.delete(document)
        db.commit()

        return {
            "message": "Document deleted successfully",
            "document_id": document_id
        }

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error deleting document {document_id}: {e}")
        raise HTTPException(status_code=500, detail=f"Failed to delete document: {str(e)}")


@router.get("/{document_id}/chunks/{chunk_index}")
async def get_document_chunk(
        document_id: int,
        chunk_index: int,
        db: Session = Depends(get_db)
):
    """Get a specific chunk from a document"""
    try:
        # Check if document exists
        document = db.query(Document).filter(Document.id == document_id).first()
        if not document:
            raise HTTPException(status_code=404, detail="Document not found")

        # Get the specific chunk
        chunk = db.query(DocumentChunk).filter(
            DocumentChunk.document_id == document_id,
            DocumentChunk.chunk_index == chunk_index
        ).first()

        if not chunk:
            raise HTTPException(status_code=404, detail="Chunk not found")

        return {
            "document_id": document_id,
            "chunk_index": chunk_index,
            "chunk_id": chunk.id,
            "content": chunk.content,
            "word_count": chunk.word_count,
            "char_count": chunk.char_count,
            "page_number": chunk.page_number,
            "document_title": document.title,
            "document_filename": document.original_filename,
            "document_file_type": document.file_type,
            "created_at": chunk.created_at.isoformat() if chunk.created_at else None
        }

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error getting chunk {chunk_index} from document {document_id}: {e}")
        raise HTTPException(status_code=500, detail=f"Failed to get chunk: {str(e)}")


@router.get("/chunk/{chunk_id_string}")
async def get_document_by_chunk_id(chunk_id_string: str, db: Session = Depends(get_db)):
    """Get document content using chunk ID format (e.g., doc_1_chunk_0)"""
    try:
        logger.info(f"Processing chunk ID: {chunk_id_string}")

        # Parse chunk ID format: doc_X_chunk_Y
        import re
        match = re.match(r'doc_(\d+)_chunk_(\d+)', chunk_id_string)
        if not match:
            raise HTTPException(status_code=400, detail=f"Invalid chunk ID format: {chunk_id_string}")

        document_id = int(match.group(1))
        chunk_index = int(match.group(2))

        logger.info(f"Extracted document_id: {document_id}, chunk_index: {chunk_index}")

        # Get document
        document = db.query(Document).filter(Document.id == document_id).first()
        if not document:
            raise HTTPException(status_code=404, detail=f"Document with ID {document_id} not found")

        # Check if file exists
        if not os.path.exists(document.file_path):
            logger.warning(f"Document file not found: {document.file_path}")
            # Don't raise error, just note it in response

        # Get the specific chunk for context
        chunk = db.query(DocumentChunk).filter(
            DocumentChunk.document_id == document_id,
            DocumentChunk.chunk_index == chunk_index
        ).first()

        # Get document content (same as view endpoint)
        content = ""
        file_type = document.file_type.lower()

        try:
            if os.path.exists(document.file_path):
                if file_type == "txt":
                    with open(document.file_path, 'r', encoding='utf-8') as f:
                        content = f.read()

                elif file_type == "pdf":
                    import PyPDF2
                    with open(document.file_path, 'rb') as f:
                        pdf_reader = PyPDF2.PdfReader(f)
                        content = ""
                        for page_num, page in enumerate(pdf_reader.pages):
                            page_text = page.extract_text()
                            content += f"\n--- Page {page_num + 1} ---\n{page_text}\n"

                elif file_type in ["doc", "docx"]:
                    from docx import Document as DocxDocument
                    doc = DocxDocument(document.file_path)
                    paragraphs = []
                    for para in doc.paragraphs:
                        if para.text.strip():
                            paragraphs.append(para.text)
                    content = "\n\n".join(paragraphs)

                else:
                    # For other file types, try to get content from chunks
                    if not content:
                        chunks = db.query(DocumentChunk).filter(
                            DocumentChunk.document_id == document_id
                        ).order_by(DocumentChunk.chunk_index).all()
                        content = "\n\n".join([c.content for c in chunks])

            else:
                # File doesn't exist, get content from database chunks
                chunks = db.query(DocumentChunk).filter(
                    DocumentChunk.document_id == document_id
                ).order_by(DocumentChunk.chunk_index).all()

                if chunks:
                    content = "\n\n".join([c.content for c in chunks])
                else:
                    content = "Document content not available"

        except Exception as e:
            logger.error(f"Error extracting content from document {document_id}: {e}")
            # Fallback to database chunks
            chunks = db.query(DocumentChunk).filter(
                DocumentChunk.document_id == document_id
            ).order_by(DocumentChunk.chunk_index).all()

            if chunks:
                content = "\n\n".join([c.content for c in chunks])
            else:
                content = f"Error extracting document content: {str(e)}"

        return {
            "id": document.id,
            "document_id": document.id,
            "filename": document.original_filename,
            "title": document.title,
            "file_type": document.file_type,
            "file_size": document.file_size,
            "word_count": document.word_count,
            "total_pages": document.total_pages,
            "content": content,
            "chunk_context": {
                "chunk_index": chunk.chunk_index if chunk else chunk_index,
                "chunk_content": chunk.content if chunk else None,
                "requested_chunk_id": chunk_id_string
            },
            "created_at": document.created_at.isoformat() if document.created_at else None,
            "processed_at": document.processed_at.isoformat() if document.processed_at else None,
            "accessed_via_chunk": True
        }

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error processing chunk ID {chunk_id_string}: {e}")
        raise HTTPException(status_code=500, detail=f"Failed to process chunk ID: {str(e)}")

@router.get("/{document_id}/chunks")
async def get_document_chunks(
        document_id: int,
        skip: int = 0,
        limit: int = 50,
        db: Session = Depends(get_db)
):
    """Get chunks for a specific document"""

    try:
        # Check if document exists
        document = db.query(Document).filter(Document.id == document_id).first()
        if not document:
            raise HTTPException(status_code=404, detail="Document not found")

        # Get chunks
        query = db.query(DocumentChunk).filter(DocumentChunk.document_id == document_id)
        total = query.count()
        chunks = query.offset(skip).limit(limit).all()

        chunk_list = []
        for chunk in chunks:
            chunk_list.append({
                "id": chunk.id,
                "chunk_index": chunk.chunk_index,
                "content": chunk.content,
                "word_count": chunk.word_count,
                "char_count": chunk.char_count,
                "page_number": chunk.page_number,
                "created_at": chunk.created_at.isoformat() if chunk.created_at else None
            })

        return {
            "document_id": document_id,
            "chunks": chunk_list,
            "total": total,
            "skip": skip,
            "limit": limit
        }

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error getting chunks for document {document_id}: {e}")
        raise HTTPException(status_code=500, detail=f"Failed to get chunks: {str(e)}")


@router.get("/stats/overview")
async def get_documents_stats(db: Session = Depends(get_db)):
    """Get overview statistics for all documents"""

    try:
        # Total documents
        total_documents = db.query(Document).count()

        # Documents by type
        from sqlalchemy import func
        type_stats = db.query(
            Document.file_type,
            func.count(Document.id).label('count')
        ).group_by(Document.file_type).all()

        # Documents by status
        status_stats = db.query(
            Document.status,
            func.count(Document.id).label('count')
        ).group_by(Document.status).all()

        # Processing status
        processing_stats = db.query(
            Document.processing_status,
            func.count(Document.id).label('count')
        ).group_by(Document.processing_status).all()

        # Total storage used
        total_size = db.query(func.sum(Document.file_size)).scalar() or 0

        # Total chunks
        total_chunks = db.query(DocumentChunk).count()

        # Recent uploads (last 7 days)
        from datetime import datetime, timedelta
        week_ago = datetime.utcnow() - timedelta(days=7)
        recent_uploads = db.query(Document).filter(Document.created_at >= week_ago).count()

        return {
            "total_documents": total_documents,
            "total_storage_bytes": total_size,
            "total_chunks": total_chunks,
            "recent_uploads_7days": recent_uploads,
            "by_file_type": {item.file_type: item.count for item in type_stats},
            "by_status": {item.status: item.count for item in status_stats},
            "by_processing_status": {item.processing_status: item.count for item in processing_stats},
            "supported_types": settings.allowed_file_types
        }

    except Exception as e:
        logger.error(f"Error getting document stats: {e}")
        raise HTTPException(status_code=500, detail=f"Failed to get stats: {str(e)}")


@router.get("/{document_id}/view")
async def view_document(document_id: str, db: Session = Depends(get_db)):  # Changed to str to handle chunk IDs
    """Get document content for viewing - supports both document IDs and chunk IDs"""

    try:
        # Check if this is a chunk ID format
        if isinstance(document_id, str) and 'chunk' in document_id:
            # Redirect to chunk handler
            return await get_document_by_chunk_id(document_id, db)

        # Convert to int for regular document ID
        try:
            doc_id = int(document_id)
        except ValueError:
            raise HTTPException(status_code=400, detail=f"Invalid document ID: {document_id}")

        document = db.query(Document).filter(Document.id == doc_id).first()

        if not document:
            raise HTTPException(status_code=404, detail="Document not found")

        if not os.path.exists(document.file_path):
            raise HTTPException(status_code=404, detail="Document file not found")

        # Get document content based on file type
        content = ""
        file_type = document.file_type.lower()

        try:
            if file_type == "txt":
                with open(document.file_path, 'r', encoding='utf-8') as f:
                    content = f.read()

            elif file_type == "pdf":
                # Extract text from PDF
                import PyPDF2
                with open(document.file_path, 'rb') as f:
                    pdf_reader = PyPDF2.PdfReader(f)
                    content = ""
                    for page_num, page in enumerate(pdf_reader.pages):
                        page_text = page.extract_text()
                        content += f"\n--- Page {page_num + 1} ---\n{page_text}\n"

            elif file_type in ["doc", "docx"]:
                # Extract text from Word documents
                from docx import Document as DocxDocument
                doc = DocxDocument(document.file_path)
                paragraphs = []
                for para in doc.paragraphs:
                    if para.text.strip():
                        paragraphs.append(para.text)
                content = "\n\n".join(paragraphs)

            else:
                raise HTTPException(status_code=400, detail=f"File type {file_type} not supported for viewing")

        except Exception as e:
            logger.error(f"Error extracting content from document {doc_id}: {e}")
            raise HTTPException(status_code=500, detail=f"Failed to extract document content: {str(e)}")

        return {
            "id": document.id,
            "document_id": document.id,
            "filename": document.original_filename,
            "title": document.title,
            "file_type": document.file_type,
            "file_size": document.file_size,
            "word_count": document.word_count,
            "total_pages": document.total_pages,
            "content": content,
            "created_at": document.created_at.isoformat() if document.created_at else None,
            "processed_at": document.processed_at.isoformat() if document.processed_at else None,
            "accessed_via_chunk": False
        }

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error viewing document {document_id}: {e}")
        raise HTTPException(status_code=500, detail=f"Failed to view document: {str(e)}")


@router.get("/{document_id}/download")
async def download_document(document_id: int, db: Session = Depends(get_db)):
    """Download original document file"""
    from fastapi.responses import FileResponse

    try:
        document = db.query(Document).filter(Document.id == document_id).first()

        if not document:
            raise HTTPException(status_code=404, detail="Document not found")

        if not os.path.exists(document.file_path):
            raise HTTPException(status_code=404, detail="Document file not found")

        # Use the current filename (which may be renamed) for the file path
        # but offer both original and current filename for download
        download_filename = document.filename  # This is the current (possibly renamed) filename

        # If user wants original filename, they can be given that option
        # For now, we'll use the current filename as it's what's actually stored

        return FileResponse(
            path=document.file_path,
            filename=download_filename,  # Use the current filename
            media_type='application/octet-stream',
            headers={
                "Content-Disposition": f'attachment; filename="{download_filename}"',
                "X-Original-Filename": document.original_filename  # Provide original name in header
            }
        )

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error downloading document {document_id}: {e}")
        raise HTTPException(status_code=500, detail=f"Failed to download document: {str(e)}")


@router.get("/{document_id}/pdf")
async def view_pdf(document_id: int, db: Session = Depends(get_db)):
    """Serve PDF file directly for viewing in browser"""
    try:
        document = db.query(Document).filter(Document.id == document_id).first()

        if not document:
            raise HTTPException(status_code=404, detail="Document not found")

        if not os.path.exists(document.file_path):
            raise HTTPException(status_code=404, detail="Document file not found")

        # Only serve PDF files through this endpoint
        if document.file_type.lower() != 'pdf':
            raise HTTPException(status_code=400, detail="This endpoint only serves PDF files")

        return FileResponse(
            path=document.file_path,
            filename=document.filename,
            media_type='application/pdf',
            headers={
                "Content-Disposition": "inline",
                "Cache-Control": "public, max-age=3600",
                "X-Original-Filename": document.original_filename
            }
        )

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error serving PDF {document_id}: {e}")
        raise HTTPException(status_code=500, detail=f"Failed to serve PDF: {str(e)}")

# Fixed background processing function
async def process_document_background(document_id: int):
    """Background task to process a document"""
    # Import here to avoid circular imports
    from app.core.database import SessionLocal

    db = None
    try:
        # Create database session
        db = SessionLocal()

        # Get document
        document = db.query(Document).filter(Document.id == document_id).first()
        if not document:
            logger.error(f"Document {document_id} not found for processing")
            return

        logger.info(
            f"Starting background processing for document {document_id}: {document.original_filename}")

        # Update status
        document.processing_status = "processing"
        db.commit()

        # Process the document
        chunks, metadata = document_processor.process_document(
            document.file_path,
            chunk_size=1000,
            chunk_overlap=200
        )

        logger.info(f"Document processed: {len(chunks)} chunks created")

        # Update document with metadata
        document.file_hash = metadata.get("file_hash")
        document.word_count = metadata.get("word_count", 0)
        document.char_count = sum(chunk.get("char_count", 0) for chunk in chunks)
        document.total_chunks = len(chunks)
        document.keywords = metadata.get("keywords", [])
        document.metadata = metadata

        # Extract additional metadata based on file type
        if document.file_type == "pdf":
            document.total_pages = metadata.get("total_pages")
        elif document.file_type in ["docx", "doc"]:
            document.total_pages = metadata.get("paragraph_count")

        # Set author, title, etc. if available in metadata
        if not document.title or document.title == document.original_filename:
            document.title = metadata.get("title") or document.original_filename

        if not document.author:
            document.author = metadata.get("author")

        if not document.subject:
            document.subject = metadata.get("subject")

        # Save chunks to database
        for chunk_data in chunks:
            chunk = DocumentChunk(
                document_id=document_id,
                chunk_index=chunk_data["chunk_index"],
                content=chunk_data["content"],
                word_count=chunk_data["word_count"],
                char_count=chunk_data["char_count"],
                page_number=chunk_data.get("page_number")
            )
            db.add(chunk)

        # Initialize ChromaDB if not already done
        await chroma_service.initialize()

        # Add to vector store
        documents_for_vector = []
        metadatas_for_vector = []
        ids_for_vector = []

        for i, chunk_data in enumerate(chunks):
            documents_for_vector.append(chunk_data["content"])
            metadatas_for_vector.append({
                "document_id": document_id,
                "chunk_index": i,
                "file_type": document.file_type,
                "filename": document.original_filename,
                "title": document.title or "",
                "category": document.category or "",
                "page_number": chunk_data.get("page_number", 0)
            })
            ids_for_vector.append(f"doc_{document_id}_chunk_{i}")

        # Add to ChromaDB
        success = await chroma_service.add_documents(
            documents=documents_for_vector,
            metadatas=metadatas_for_vector,
            ids=ids_for_vector
        )

        if success:
            # Update document status
            document.status = "completed"
            document.processing_status = "completed"
            document.is_processed = True
            document.is_searchable = True
            document.processed_at = func.now()
            document.error_message = None

            logger.info(f"Document {document_id} processed successfully")
        else:
            raise Exception("Failed to add documents to vector store")

        db.commit()

    except Exception as e:
        logger.error(f"Error processing document {document_id}: {e}")

        if db:
            try:
                # Update document with error status
                document = db.query(Document).filter(Document.id == document_id).first()
                if document:
                    document.status = "failed"
                    document.processing_status = "failed"
                    document.error_message = str(e)
                    db.commit()
            except Exception as commit_error:
                logger.error(f"Error updating document status: {commit_error}")

    finally:
        if db:
            db.close()

    # Additional utility functions for document processing
    def get_file_extension(filename: str) -> str:
        """Get file extension from filename"""
        return filename.lower().split('.')[-1] if '.' in filename else ''

    def generate_unique_filename(original_filename: str) -> str:
        """Generate unique filename while preserving extension"""
        import uuid
        file_extension = get_file_extension(original_filename)
        return f"{uuid.uuid4()}.{file_extension}"

    def calculate_file_hash(file_path: str) -> str:
        """Calculate SHA-256 hash of file"""
        import hashlib

        hash_sha256 = hashlib.sha256()
        try:
            with open(file_path, "rb") as f:
                for chunk in iter(lambda: f.read(4096), b""):
                    hash_sha256.update(chunk)
            return hash_sha256.hexdigest()
        except Exception as e:
            logger.error(f"Error calculating file hash: {e}")
            return ""

    # Batch processing endpoints
@router.post("/batch/upload")
async def batch_upload_documents(
        background_tasks: BackgroundTasks,
        files: List[UploadFile] = File(...),
        category: Optional[str] = Form(None),
        description: Optional[str] = Form(None),
        auto_process: bool = Form(True),
        db: Session = Depends(get_db)
):
    """Upload multiple documents at once"""

    if len(files) > 50:  # Limit batch size
        raise HTTPException(status_code=400, detail="Too many files. Maximum 50 files per batch.")

    results = []
    upload_dir = "uploads"
    os.makedirs(upload_dir, exist_ok=True)

    for file in files:
        try:
            # Validate file
            if not validate_file_type(file.filename):
                results.append({
                    "filename": file.filename,
                    "status": "failed",
                    "error": f"Unsupported file type"
                })
                continue

            # Read and validate size
            file_content = await file.read()
            if not validate_file_size(len(file_content)):
                results.append({
                    "filename": file.filename,
                    "status": "failed",
                    "error": "File too large"
                })
                continue

            # Save file
            unique_filename = generate_unique_filename(file.filename)
            file_path = os.path.join(upload_dir, unique_filename)

            with open(file_path, "wb") as buffer:
                buffer.write(file_content)

            # Get file type
            file_type = document_processor.get_file_type(file_path)

            # Create database record
            document = Document(
                filename=unique_filename,
                original_filename=file.filename,
                file_path=file_path,
                file_type=file_type,
                file_size=len(file_content),
                title=file.filename,
                category=category,
                description=description,
                status="uploaded",
                processing_status="pending"
            )

            db.add(document)
            db.commit()
            db.refresh(document)

            # Schedule processing
            if auto_process:
                background_tasks.add_task(process_document_background, document.id)

            results.append({
                "filename": file.filename,
                "document_id": document.id,
                "status": "uploaded",
                "file_type": file_type,
                "size": len(file_content)
            })

            logger.info(f"Batch uploaded: {file.filename} -> {document.id}")

        except Exception as e:
            logger.error(f"Error uploading {file.filename}: {e}")
            results.append({
                "filename": file.filename,
                "status": "failed",
                "error": str(e)
            })

    successful_uploads = len([r for r in results if r["status"] == "uploaded"])

    return {
        "message": f"Batch upload completed: {successful_uploads}/{len(files)} files uploaded",
        "results": results,
        "total_files": len(files),
        "successful_uploads": successful_uploads,
        "auto_process": auto_process
    }

@router.post("/batch/process")
async def batch_process_documents(
        background_tasks: BackgroundTasks,
        document_ids: List[int],
        db: Session = Depends(get_db)
):
    """Process multiple documents"""

    if len(document_ids) > 20:  # Limit batch processing
        raise HTTPException(status_code=400, detail="Too many documents. Maximum 20 documents per batch.")

    results = []

    for doc_id in document_ids:
        try:
            document = db.query(Document).filter(Document.id == doc_id).first()

            if not document:
                results.append({
                    "document_id": doc_id,
                    "status": "failed",
                    "error": "Document not found"
                })
                continue

            if document.processing_status == "processing":
                results.append({
                    "document_id": doc_id,
                    "status": "skipped",
                    "error": "Already processing"
                })
                continue

            # Update status and schedule processing
            document.processing_status = "processing"
            db.commit()

            background_tasks.add_task(process_document_background, doc_id)

            results.append({
                "document_id": doc_id,
                "filename": document.original_filename,
                "status": "scheduled"
            })

        except Exception as e:
            logger.error(f"Error scheduling processing for document {doc_id}: {e}")
            results.append({
                "document_id": doc_id,
                "status": "failed",
                "error": str(e)
            })

    scheduled_count = len([r for r in results if r["status"] == "scheduled"])

    return {
        "message": f"Batch processing scheduled: {scheduled_count}/{len(document_ids)} documents",
        "results": results,
        "total_documents": len(document_ids),
        "scheduled_processing": scheduled_count
    }



# Health check endpoint for documents service
@router.get("/health")
async def documents_health_check(db: Session = Depends(get_db)):
    """Health check for documents service"""
    try:
        # Test database connection
        total_docs = db.query(Document).count()

        # Test document processor
        processor_status = "ok" if document_processor else "error"

        # Test chroma service
        chroma_status = "ok"
        try:
            await chroma_service.initialize()
        except Exception:
            chroma_status = "error"

        return {
            "status": "healthy",
            "database": "connected",
            "total_documents": total_docs,
            "document_processor": processor_status,
            "vector_store": chroma_status,
            "supported_types": settings.allowed_file_types
        }

    except Exception as e:
        logger.error(f"Health check failed: {e}")
        raise HTTPException(status_code=503, detail=f"Service unhealthy: {str(e)}")


@router.get("/{document_id}/view-as-pdf")
async def view_document_as_pdf(document_id: int, db: Session = Depends(get_db)):
    """Convert document to PDF for viewing if it's not already a PDF"""
    try:
        document = db.query(Document).filter(Document.id == document_id).first()

        if not document:
            raise HTTPException(status_code=404, detail="Document not found")

        if not os.path.exists(document.file_path):
            raise HTTPException(status_code=404, detail="Document file not found")

        # If it's already a PDF, serve it directly
        if document.file_type.lower() == 'pdf':
            return FileResponse(
                path=document.file_path,
                filename=document.filename,
                media_type='application/pdf',
                headers={
                    "Content-Disposition": "inline",
                    "Cache-Control": "public, max-age=3600"
                }
            )

        # Check if document type can be converted
        if not document_processor.can_convert_to_pdf(document.file_type):
            raise HTTPException(
                status_code=400,
                detail=f"Cannot convert {document.file_type} files to PDF for viewing"
            )

        # Convert to PDF
        try:
            pdf_path = document_processor.convert_to_pdf_for_viewing(document.file_path)

            return FileResponse(
                path=pdf_path,
                filename=f"{os.path.splitext(document.filename)[0]}.pdf",
                media_type='application/pdf',
                headers={
                    "Content-Disposition": "inline",
                    "Cache-Control": "public, max-age=3600",
                    "X-Converted-From": document.file_type,
                    "X-Original-Filename": document.filename
                }
            )

        except Exception as conversion_error:
            logger.error(f"PDF conversion failed for document {document_id}: {conversion_error}")
            raise HTTPException(
                status_code=500,
                detail=f"Failed to convert document to PDF: {str(conversion_error)}"
            )

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error serving document as PDF {document_id}: {e}")
        raise HTTPException(status_code=500, detail=f"Failed to serve document as PDF: {str(e)}")

@router.get("/{document_id}/content")
async def get_document_content(document_id: int, db: Session = Depends(get_db)):
    """Get document content for viewing"""
    try:
        document = db.query(Document).filter(Document.id == document_id).first()
        if not document:
            raise HTTPException(status_code=404, detail="Document not found")

        # Get chunks from database
        chunks = db.query(DocumentChunk).filter(
            DocumentChunk.document_id == document_id
        ).order_by(DocumentChunk.chunk_index).all()

        if not chunks:
            raise HTTPException(status_code=404, detail="Document content not found")

        # Format content
        content = {
            "document_id": document.id,
            "filename": document.original_filename,
            "title": document.title,
            "file_type": document.file_type,
            "total_chunks": len(chunks),
            "content": "\n\n".join([chunk.content for chunk in chunks]),
            "chunks": [
                {
                    "index": chunk.chunk_index,
                    "content": chunk.content,
                    "page_number": chunk.page_number,
                    "word_count": chunk.word_count
                }
                for chunk in chunks
            ],
            "metadata": {
                "file_size": document.file_size,
                "word_count": document.word_count,
                "total_pages": document.total_pages,
                "category": document.category,
                "author": document.author,
                "created_at": document.created_at.isoformat() if document.created_at else None
            }
        }

        return content

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error getting document content {document_id}: {e}")
        raise HTTPException(status_code=500, detail=f"Failed to get document content: {str(e)}")